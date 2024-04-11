"""File upload handling and text extraction."""
import os
import uuid
from pathlib import Path

from fastapi import UploadFile

from app.core.config import settings

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}
MAX_FILE_SIZE = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024


def validate_file(file: UploadFile) -> bool:
    ext = Path(file.filename).suffix.lower()
    return ext in ALLOWED_EXTENSIONS


async def save_upload(file: UploadFile, subdir: str = "resumes", max_size_mb: int | None = None) -> str:
    ext = Path(file.filename).suffix.lower()
    unique_name = f"{uuid.uuid4().hex}{ext}"
    upload_dir = os.path.join(settings.UPLOAD_DIR, subdir)
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, unique_name)

    limit_mb = max_size_mb or settings.MAX_UPLOAD_SIZE_MB
    max_bytes = limit_mb * 1024 * 1024
    content = await file.read()
    if len(content) > max_bytes:
        raise ValueError(f"File size exceeds {limit_mb}MB limit")

    with open(file_path, "wb") as f:
        f.write(content)

    return file_path


def extract_text_from_pdf(file_path: str) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception:
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Extract table text (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    parts.append(" | ".join(row_text))

        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_resume_text(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    return ""
